from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """Set background shading for a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_row_height(row, height):
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'atLeast')
    row._tr.get_or_add_trPr().append(trHeight)

def set_table_borders(table):
    """Set standard borders for a table."""
    tbl = table._tbl
    tblPr = tbl.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '4')
        node.set(qn('w:color'), 'cccccc')
        tblBorders.append(node)
    tblPr.append(tblBorders)

def create_improved_capex(output_path):
    doc = Document()
    
    # Set narrow margins for 1-page fit
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(7.5)
    
    # Define Colors
    NAVY = RGBColor(0x1e, 0x3a, 0x8a)
    GREY_BG = 'F3F4F6'
    
    # 1. Header Branding
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(7.3)
    
    # Logo in the left cell
    logo_cell = header_table.rows[0].cells[0]
    try:
        run_logo = logo_cell.paragraphs[0].add_run()
        run_logo.add_picture('berkeley_logo.jpg', width=Inches(1.8))
    except Exception as e:
        logo_cell.text = "Berkeley Logo"
        print(f"Logo error: {e}")
    
    # Portal text in the right cell
    portal_cell = header_table.rows[0].cells[1]
    p_portal = portal_cell.paragraphs[0]
    p_portal.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_portal.add_run('eSign Document Portal')
    run.font.size = Pt(7.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
    
    # Title
    title = doc.add_paragraph('CAPITAL EXPENDITURE REQUEST')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    
    # Helper to add section title
    def add_section_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        # Add a bottom border to the paragraph
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '1e3a8a')
        pbdr.append(bottom)
        p._p.get_or_add_pPr().append(pbdr)

    # --- Section: Request Info ---
    add_section_title('General Information')
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    table.columns[0].width = Inches(1.4)
    table.columns[1].width = Inches(2.2)
    table.columns[2].width = Inches(1.4)
    table.columns[3].width = Inches(2.3)
    
    def fill_cell(cell, label, value):
        p = cell.paragraphs[0]
        run_l = p.add_run(f'{label}: ')
        run_l.bold = True
        run_l.font.size = Pt(7.5)
        run_v = p.add_run(value)
        run_v.font.size = Pt(7.5)
        # Add spacing
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    fill_cell(table.rows[0].cells[0], 'Division', '{{ requesting_division }}')
    fill_cell(table.rows[0].cells[2], 'Branch', '{{ branch }}')
    fill_cell(table.rows[1].cells[0], 'Date', '{{ request_date }}')
    fill_cell(table.rows[1].cells[2], 'Staff', '{{ staff_name }}')
    
    # --- Section: Location ---
    add_section_title('Location & Item Status')
    table2 = doc.add_table(rows=2, cols=4)
    table2.autofit = False
    table2.columns[0].width = Inches(1.4)
    table2.columns[1].width = Inches(2.2)
    table2.columns[2].width = Inches(1.4)
    table2.columns[3].width = Inches(2.3)
    
    fill_cell(table2.rows[0].cells[0], 'Location', '{{ location }}')
    fill_cell(table2.rows[0].cells[2], 'Site Code', '{{ site_code }}')
    fill_cell(table2.rows[1].cells[0], 'New Item?', '{{ is_new_item_yes_no }}')
    fill_cell(table2.rows[1].cells[2], 'Replacement?', '{{ is_replacement_yes_no }}')

    # --- Section: Expenditure Details ---
    add_section_title('Expenditure Items')
    item_table = doc.add_table(rows=1, cols=5)
    item_table.autofit = False
    item_table.style = 'Table Grid'
    
    # Set Logical Widths
    # Total ~7.3 inches
    item_table.columns[0].width = Inches(0.4)  # #
    item_table.columns[1].width = Inches(3.8)  # Item Description
    item_table.columns[2].width = Inches(0.7)  # Budgeted
    item_table.columns[3].width = Inches(1.2)  # Date Required
    item_table.columns[4].width = Inches(1.2)  # Amount (AED)

    hdr = item_table.rows[0].cells
    labels = ['#', 'Item Description', 'Budgeted', 'Date Required', 'Amount (AED)']
    for i, label in enumerate(labels):
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(7.5)
        set_cell_shading(hdr[i], GREY_BG)
        
    for i in range(1, 6):
        row_cells = item_table.add_row().cells
        # Increase row height ~ 0.25 inch = 360 twips (1440 twips = 1 inch)
        # or just rely on padding. Let's use paragraph spacing.
        row_cells[0].text = str(i)
        row_cells[1].text = f'{{{{ item_{i}_description }}}}'
        row_cells[2].text = f'{{{{ is_item_{i}_budgeted_yes_no }}}}'
        row_cells[3].text = f'{{{{ item_{i}_date_required }}}}'
        row_cells[4].text = f'{{{{ item_{i}_amount }}}}'
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set font size and spacing for cells
        for idx, cell in enumerate(row_cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                # Specific horizontal alignment
                if idx == 0: # Index
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif idx == 1: # Description
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif idx >= 2 and idx != 4: # Budgeted, Date
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Amount is already RIGHT
                
                for run in para.runs:
                    run.font.size = Pt(7.5)

    # --- Justification ---
    add_section_title('Justification')
    p = doc.add_paragraph('{{ justification }}')
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(7.5)

    # --- Approvals ---
    add_section_title('Approval Workflow')
    sig_table = doc.add_table(rows=0, cols=4)
    sig_table.style = 'Table Grid'
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(0.5)
    sig_table.columns[1].width = Inches(3.5)
    sig_table.columns[2].width = Inches(2.0)
    sig_table.columns[3].width = Inches(1.3)
    
    header = sig_table.add_row().cells
    header[0].text = 'Step'
    header[1].text = 'Approver Details'
    header[2].text = 'Signature'
    header[3].text = 'Date Signed'
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(7.5)
        set_cell_shading(cell, GREY_BG)

    for i in range(1, 6):
        row = sig_table.add_row().cells
        # Set height for signature row to give space
        set_row_height(sig_table.rows[i], 700) # ~0.5 inch

        # Col 0: Step
        row[0].text = f'{i}'
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Col 1: Approver Details
        info = row[1].add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        name_run = info.add_run(f'{{{{ approver_{i}_name }}}}')
        name_run.bold = True
        name_run.font.size = Pt(7.5)
        pos_run = info.add_run(f'\n{{{{ approver_{i}_position }}}}')
        pos_run.font.size = Pt(7)
        
        # Col 2: Signature
        row[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        # Add the line
        sig_p = row[2].add_paragraph()
        sig_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_sig = sig_p.add_run('________________________')
        run_sig.font.size = Pt(7)

        # Col 3: Date
        row[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        date_p = row[3].add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = date_p.add_run('__________')
        run_date.font.size = Pt(7)

    doc.save(output_path)
    print(f"Created professional template with logo and reduced font at: {output_path}")

if __name__ == "__main__":
    create_improved_capex("Capex_Template_New.docx")
