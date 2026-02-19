from docx import Document
import sys

def modify_template(input_path, output_path):
    doc = Document(input_path)
    
    # 1. General replacements in paragraphs
    for para in doc.paragraphs:
        if "Dubai" in para.text:
            para.text = para.text.replace("Dubai", "{{ location }}")
        if "Head Office" in para.text:
            para.text = para.text.replace("Head Office", "{{ site }}")

    # 2. Table replacements
    for table in doc.tables:
        for row in table.rows:
            # Helper to get full text of row for logic check
            row_text = "".join([c.text for c in row.cells])
            
            # Replace common fields
            for cell in row.cells:
                if "Dubai" in cell.text:
                    cell.text = cell.text.replace("Dubai", "{{ location }}")
                if "Head Office" in cell.text:
                    cell.text = cell.text.replace("Head Office", "{{ site }}")
                # Replace date placeholders
                if "….… /…… /………" in cell.text:
                    cell.text = cell.text.replace("….… /…… /………", "{{ date }}")

            # Signature Anchors
            # Logic: If row contains a specific name, replace (Signature): with a specific Anchor
            
            if "Muhammad Bilal" in row_text:
                for cell in row.cells:
                    if "(Signature):" in cell.text:
                        cell.text = "[SIGN_PROCUREMENT]"
            
            elif "Lennart" in row_text:
                for cell in row.cells:
                    if "(Signature):" in cell.text:
                        cell.text = "[SIGN_IT]"
            
            elif "Jawad Ahmad Bhatti" in row_text:
                 for cell in row.cells:
                    if "(Signature):" in cell.text:
                        cell.text = "[SIGN_CFO]"
            
            elif "Karl-Heinz Mair" in row_text:
                 for cell in row.cells:
                    if "(Signature):" in cell.text:
                        cell.text = "[SIGN_CEO]"

    doc.save(output_path)
    print(f"Saved modified template to {output_path}")

if __name__ == "__main__":
    modify_template("Laptop Mair.docx", "Laptop_Mair_Template.docx")
