from docx import Document
import sys

def modify_capex(input_path, output_path):
    doc = Document(input_path)
    
    # 1. General replacements (Appending placeholders to labels)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if "Reason (Justification) for Purchase:" in text and "{{ reason }}" not in text:
                    cell.text = text + " {{ reason }}"
                elif "Where used (Location):" in text and "{{ location }}" not in text:
                     cell.text = text + " {{ location }}"
                elif "Site Code:" in text and "{{ site_code }}" not in text:
                     cell.text = text + " {{ site_code }}"
                elif "Total Value" in text and "{{ total_value }}" not in text:
                     cell.text = text.replace("Total Value", "Total Value {{ total_value }}")
                elif "Date Expenditure" in text and "{{ date }}" not in text:
                     cell.text = text + " {{ date }}"
    
    # 2. Approval Rows
    approver_count = 0
    for table in doc.tables:
        is_approval_table = False
        for row in table.rows:
            if "Approved By:" in row.cells[0].text:
                is_approval_table = True
                break
        
        if is_approval_table:
            # First pass: Identify all relevant rows to count them correctly if needed
            # But simple iteration worked before, let's stick to it but refining the conditions
            
            for row in table.rows:
                cells = row.cells
                row_text = "".join([c.text for c in cells])
                
                # Check column 0 for 'Approved By:' and any column for the placeholder text
                # We specifically look for the row that has the Name/Position placeholder we want to replace
                if "Approved By:" in cells[0].text and ("(Name) (Position)" in row_text or "{{ approver_" in row_text):
                    # Note: We check {{ approver_ }} to avoid re-modifying if run multiple times, 
                    # but since we start from fresh input it's fine.
                    # Wait, if we run on fresh input, it will have (Name) (Position).
                    
                    if "(Name) (Position)" in row_text: 
                        approver_count += 1
                        
                        for cell in cells:
                            if "(Name) (Position)" in cell.text:
                                cell.text = f"{{{{ approver_{approver_count}_name }}}}\n{{{{ approver_{approver_count}_position }}}}"
                            
                            if "(Signature):" in cell.text:
                                # Conditional Anchor: Only show [SIGN_X] if approver_X_name is present
                                cell.text = f"{{% if approver_{approver_count}_name %}}[SIGN_{approver_count}]{{% endif %}}"

    doc.save(output_path)
    print(f"Saved modified Capex template to {output_path}. Found {approver_count} approval rows.")

if __name__ == "__main__":
    modify_capex("Capex Template.docx", "Capex_Template_Ready.docx")
