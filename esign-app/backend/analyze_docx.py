import re
from docx import Document
import sys

def analyze_docx(file_path):
    try:
        doc = Document(file_path)
        print(f"--- Document Content ({file_path}) ---")
        
        placeholders = set()
        pattern = re.compile(r'\{\{(.*?)\}\}')
        
        for para in doc.paragraphs:
            matches = pattern.findall(para.text)
            for m in matches:
                placeholders.add(m.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = pattern.findall(cell.text)
                    for m in matches:
                        placeholders.add(m.strip())
        
        print("\n--- Placeholders Found ---")
        for p in sorted(list(placeholders)):
            print(f"  {{{{ {p} }}}}")
            
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        analyze_docx(sys.argv[1])
    else:
        print("Usage: python analyze_docx.py <path_to_docx>")
